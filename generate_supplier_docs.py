from fastapi import UploadFile, File, APIRouter, HTTPException
from fastapi.responses import FileResponse
import pandas as pd
from docx import Document
from io import BytesIO
import zipfile
import os

router = APIRouter()

SUPPLIER_MAP_PATH = "supplier.csv.csv"
TEMP_DIR = "/mnt/data/supplier_docs"
ZIP_PATH = os.path.join(TEMP_DIR, "supplier_outputs.zip")

@router.post("/generate_supplier_docs")
async def generate_supplier_docs(file: UploadFile = File(...)):
    try:
        # 1. Load order Excel file
        order_bytes = await file.read()
        order_df = pd.read_excel(BytesIO(order_bytes))

        # 2. Normalize column names
        COLUMN_ALIASES = {
            "ORDER NO": "ORDER",
            "ORDER NUMBER": "ORDER",
            "ORDER#": "ORDER",
            "PRODUCT CODE": "SKU",
            "ITEM CODE": "SKU",
            "OFFER SKU": "SKU",
            "QUANTITY": "QTY",
            "QTY.": "QTY",
            "QTY ORDERED": "QTY"
        }
        order_df.columns = [COLUMN_ALIASES.get(c.strip().upper(), c.strip().upper()) for c in order_df.columns]

        if not {"SKU", "QTY"}.issubset(order_df.columns):
            raise HTTPException(status_code=400, detail="Missing required columns SKU or QTY")

        orders = order_df[["SKU", "QTY"]].dropna().to_dict(orient="records")

        # 3. Load supplier map
        supplier_df = pd.read_csv(SUPPLIER_MAP_PATH)
        supplier_lookup = dict(zip(supplier_df['Supplier SKU'], supplier_df['Supplier Name']))

        # 4. Split orders
        supplier_orders = {}
        unmatched = []
        for entry in orders:
            sku = str(entry['SKU']).strip()
            qty = int(entry['QTY'])
            supplier = supplier_lookup.get(sku)
            if supplier:
                supplier_orders.setdefault(supplier, []).append({"SKU": sku, "QTY": qty})
            else:
                unmatched.append({"SKU": sku, "QTY": qty})

        # 5. Generate output docs
        os.makedirs(TEMP_DIR, exist_ok=True)
        output_files = []

        for supplier, items in supplier_orders.items():
            doc = Document()
            doc.add_heading(f"{supplier} Order", level=1)
            for item in items:
                doc.add_paragraph(f"SKU: {item['SKU']} — QTY: {item['QTY']}")
            path = os.path.join(TEMP_DIR, f"{supplier}_Orders.docx")
            doc.save(path)
            output_files.append(path)

        if supplier_orders.get("Nisbets"):
            checklist_path = os.path.join(TEMP_DIR, "Nisbets_Checklist.csv")
            pd.DataFrame(supplier_orders["Nisbets"]).to_csv(checklist_path, index=False)
            output_files.append(checklist_path)

        # 6. Zip all
        with zipfile.ZipFile(ZIP_PATH, "w") as zipf:
            for f in output_files:
                zipf.write(f, os.path.basename(f))

        return {
            "zip_file": ZIP_PATH,
            "unmatched_skus": unmatched
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to process: {str(e)}")
