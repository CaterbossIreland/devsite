from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import StreamingResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
import pandas as pd
from io import BytesIO
from docx import Document
import csv
from pydantic import BaseModel
import os
import requests

# Graph API authentication
TENANT_ID = "ce280aae-ee92-41fe-ab60-66b37ebc97dd"
CLIENT_ID = "83acd574-ab02-4cfe-b28c-e38c733d9a52"
CLIENT_SECRET = "FYX8Q~bZVXuKEenMTryxYw-ZuQOq2OBTNIu8Qa~i"
DRIVE_ID = "b!_osuAVwo5EyWvEWEMnzopleQal6puNREsmylMfjWpjsv-rD7sQmrQLHDhsQKjaxA"
SUPPLIER_PATH = "/Supplier.csv"

def get_access_token():
    url = f"https://login.microsoftonline.com/{TENANT_ID}/oauth2/v2.0/token"
    headers = {"Content-Type": "application/x-www-form-urlencoded"}
    data = {
        "client_id": CLIENT_ID,
        "scope": "https://graph.microsoft.com/.default",
        "client_secret": CLIENT_SECRET,
        "grant_type": "client_credentials"
    }
    response = requests.post(url, headers=headers, data=data)
    return response.json().get("access_token")

def download_supplier_csv():
    access_token = get_access_token()
    headers = {"Authorization": f"Bearer {access_token}"}
    url = f"https://graph.microsoft.com/v1.0/drives/{DRIVE_ID}/root:{SUPPLIER_PATH}:/content"
    response = requests.get(url, headers=headers)
    if response.status_code == 200:
        return pd.read_csv(BytesIO(response.content))
    else:
        raise HTTPException(status_code=500, detail="Failed to download supplier.csv")

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

class OrderItem(BaseModel):
    SKU: str
    QTY: int

class OrderData(BaseModel):
    rows: list[OrderItem]

@app.post("/generate-docs/")
async def generate_docs(file: UploadFile = File(...)):
    try:
        # Load the order file
        order_df = pd.read_excel(file.file)
        order_df["Offer SKU"] = order_df["Offer SKU"].astype(str).str.strip()

        # Load supplier mapping
        supplier_df = download_supplier_csv()
        supplier_df["SKU"] = supplier_df["SKU"].astype(str).str.strip()
        supplier_df["Supplier"] = supplier_df["Supplier"].astype(str).str.strip()

        # Merge
        merged = pd.merge(order_df, supplier_df, how="left", left_on="Offer SKU", right_on="SKU")

        # Check for unmatched SKUs
        unmatched = merged[merged["Supplier"].isnull()]
        if not unmatched.empty:
            unmatched_list = unmatched["Offer SKU"].dropna().unique().tolist()
            return JSONResponse(status_code=400, content={"unmatched_skus": unmatched_list})

        # Group by supplier
        supplier_docs = {}
        supplier_csvs = {}

        for supplier, group in merged.groupby("Supplier"):
            # Word doc
            doc = Document()
            doc.add_heading(f"Order for {supplier}", 0)
            for idx, row in group.iterrows():
                doc.add_paragraph(f"{row['Offer SKU']} - {row.get('Description', '')} x {row.get('Qty', '')}")
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            supplier_docs[supplier] = buffer

            # CSV for Nisbets only
            if supplier.lower() == "nisbets":
                nisbets_buffer = BytesIO()
                writer = csv.writer(nisbets_buffer)
                writer.writerow(["SKU", "QTY"])
                for idx, row in group.iterrows():
                    writer.writerow([row["Offer SKU"], row.get("Qty", "")])
                nisbets_buffer.seek(0)
                supplier_csvs[supplier] = nisbets_buffer

        # Create zip
        zip_buffer = BytesIO()
        import zipfile
        with zipfile.ZipFile(zip_buffer, "w") as zipf:
            for name, buf in supplier_docs.items():
                zipf.writestr(f"{name}_Orders.docx", buf.read())
            for name, buf in supplier_csvs.items():
                zipf.writestr(f"{name}_Checklist.csv", buf.read())
        zip_buffer.seek(0)

        return StreamingResponse(zip_buffer, media_type="application/x-zip-compressed", headers={
            "Content-Disposition": "attachment; filename=supplier_docs.zip"
        })

    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})
